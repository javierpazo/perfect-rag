"""Document loaders for various file formats."""

import hashlib
import mimetypes
from abc import ABC, abstractmethod
from datetime import datetime
from pathlib import Path
from typing import Any, BinaryIO

import structlog

from perfect_rag.models.document import Document, DocumentMetadata, DocumentStatus

logger = structlog.get_logger(__name__)


class DocumentLoader(ABC):
    """Abstract base class for document loaders."""

    supported_extensions: list[str] = []
    supported_mimetypes: list[str] = []

    @abstractmethod
    async def load(
        self,
        source: str | Path | BinaryIO,
        metadata: dict[str, Any] | None = None,
    ) -> Document:
        """Load document from source.

        Args:
            source: File path, URL, or file-like object
            metadata: Optional metadata to attach

        Returns:
            Document object with extracted content
        """
        pass

    @classmethod
    def supports(cls, path: str | Path) -> bool:
        """Check if loader supports this file type."""
        path = Path(path) if isinstance(path, str) else path
        ext = path.suffix.lower()
        return ext in cls.supported_extensions


class TextLoader(DocumentLoader):
    """Load plain text files."""

    supported_extensions = [".txt", ".md", ".rst", ".log"]
    supported_mimetypes = ["text/plain", "text/markdown"]

    async def load(
        self,
        source: str | Path | BinaryIO,
        metadata: dict[str, Any] | None = None,
    ) -> Document:
        if isinstance(source, (str, Path)):
            path = Path(source)
            with open(path, "r", encoding="utf-8") as f:
                content = f.read()
            filename = path.name
            file_size = path.stat().st_size
        else:
            content = source.read()
            if isinstance(content, bytes):
                content = content.decode("utf-8")
            filename = getattr(source, "name", "unknown.txt")
            file_size = len(content.encode("utf-8"))

        doc_id = self._generate_id(content)
        meta = metadata or {}

        return Document(
            id=doc_id,
            content=content,
            metadata=DocumentMetadata(
                title=meta.get("title", filename),
                source=str(source) if isinstance(source, (str, Path)) else filename,
                mime_type="text/plain",
                file_size=file_size,
                language=meta.get("language"),
                author=meta.get("author"),
                created_at=meta.get("created_at"),
                tags=meta.get("tags", []),
                custom=meta.get("custom", {}),
            ),
            status=DocumentStatus.PENDING,
            acl=meta.get("acl", ["*"]),
        )

    def _generate_id(self, content: str) -> str:
        """Generate document ID from content hash."""
        return hashlib.sha256(content.encode()).hexdigest()[:16]


class PDFLoader(DocumentLoader):
    """Load PDF documents."""

    supported_extensions = [".pdf"]
    supported_mimetypes = ["application/pdf"]

    async def load(
        self,
        source: str | Path | BinaryIO,
        metadata: dict[str, Any] | None = None,
    ) -> Document:
        try:
            import pymupdf as fitz
        except ImportError:
            try:
                import fitz
            except ImportError:
                raise ImportError("pymupdf required for PDF loading: pip install pymupdf")

        if isinstance(source, (str, Path)):
            path = Path(source)
            doc = fitz.open(path)
            filename = path.name
            file_size = path.stat().st_size
        else:
            data = source.read()
            doc = fitz.open(stream=data, filetype="pdf")
            filename = getattr(source, "name", "unknown.pdf")
            file_size = len(data)

        # Extract text from all pages
        text_parts = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text_parts.append(page.get_text())

        content = "\n\n".join(text_parts)
        doc.close()

        doc_id = self._generate_id(content)
        meta = metadata or {}

        return Document(
            id=doc_id,
            content=content,
            metadata=DocumentMetadata(
                title=meta.get("title", filename),
                source=str(source) if isinstance(source, (str, Path)) else filename,
                mime_type="application/pdf",
                file_size=file_size,
                page_count=len(text_parts),
                language=meta.get("language"),
                author=meta.get("author"),
                created_at=meta.get("created_at"),
                tags=meta.get("tags", []),
                custom=meta.get("custom", {}),
            ),
            status=DocumentStatus.PENDING,
            acl=meta.get("acl", ["*"]),
        )

    def _generate_id(self, content: str) -> str:
        return hashlib.sha256(content.encode()).hexdigest()[:16]


class HTMLLoader(DocumentLoader):
    """Load HTML documents."""

    supported_extensions = [".html", ".htm"]
    supported_mimetypes = ["text/html"]

    async def load(
        self,
        source: str | Path | BinaryIO,
        metadata: dict[str, Any] | None = None,
    ) -> Document:
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("beautifulsoup4 required: pip install beautifulsoup4")

        if isinstance(source, (str, Path)):
            path = Path(source)
            with open(path, "r", encoding="utf-8") as f:
                html_content = f.read()
            filename = path.name
            file_size = path.stat().st_size
        else:
            html_content = source.read()
            if isinstance(html_content, bytes):
                html_content = html_content.decode("utf-8")
            filename = getattr(source, "name", "unknown.html")
            file_size = len(html_content.encode("utf-8"))

        # Parse HTML and extract text
        soup = BeautifulSoup(html_content, "html.parser")

        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()

        # Get text
        content = soup.get_text(separator="\n", strip=True)

        # Try to extract title
        title = soup.title.string if soup.title else filename

        doc_id = self._generate_id(content)
        meta = metadata or {}

        return Document(
            id=doc_id,
            content=content,
            metadata=DocumentMetadata(
                title=meta.get("title", title),
                source=str(source) if isinstance(source, (str, Path)) else filename,
                mime_type="text/html",
                file_size=file_size,
                language=meta.get("language"),
                author=meta.get("author"),
                created_at=meta.get("created_at"),
                tags=meta.get("tags", []),
                custom=meta.get("custom", {}),
            ),
            status=DocumentStatus.PENDING,
            acl=meta.get("acl", ["*"]),
        )

    def _generate_id(self, content: str) -> str:
        return hashlib.sha256(content.encode()).hexdigest()[:16]


class DocxLoader(DocumentLoader):
    """Load Microsoft Word documents."""

    supported_extensions = [".docx"]
    supported_mimetypes = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]

    async def load(
        self,
        source: str | Path | BinaryIO,
        metadata: dict[str, Any] | None = None,
    ) -> Document:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx required: pip install python-docx")

        if isinstance(source, (str, Path)):
            path = Path(source)
            doc = DocxDocument(path)
            filename = path.name
            file_size = path.stat().st_size
        else:
            doc = DocxDocument(source)
            filename = getattr(source, "name", "unknown.docx")
            source.seek(0)
            file_size = len(source.read())

        # Extract text from paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)

        # Try to get document properties
        try:
            core_props = doc.core_properties
            title = core_props.title or filename
            author = core_props.author
            created = core_props.created
        except Exception:
            title = filename
            author = None
            created = None

        doc_id = self._generate_id(content)
        meta = metadata or {}

        return Document(
            id=doc_id,
            content=content,
            metadata=DocumentMetadata(
                title=meta.get("title", title),
                source=str(source) if isinstance(source, (str, Path)) else filename,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                file_size=file_size,
                language=meta.get("language"),
                author=meta.get("author", author),
                created_at=meta.get("created_at", created),
                tags=meta.get("tags", []),
                custom=meta.get("custom", {}),
            ),
            status=DocumentStatus.PENDING,
            acl=meta.get("acl", ["*"]),
        )

    def _generate_id(self, content: str) -> str:
        return hashlib.sha256(content.encode()).hexdigest()[:16]


class JSONLoader(DocumentLoader):
    """Load JSON documents."""

    supported_extensions = [".json", ".jsonl"]
    supported_mimetypes = ["application/json"]

    def __init__(self, content_key: str | None = None, jq_filter: str | None = None):
        """
        Args:
            content_key: Key to extract content from JSON (e.g., "text", "content")
            jq_filter: Optional jq-style filter to extract content
        """
        self.content_key = content_key
        self.jq_filter = jq_filter

    async def load(
        self,
        source: str | Path | BinaryIO,
        metadata: dict[str, Any] | None = None,
    ) -> Document:
        import json

        if isinstance(source, (str, Path)):
            path = Path(source)
            with open(path, "r", encoding="utf-8") as f:
                raw_content = f.read()
            filename = path.name
            file_size = path.stat().st_size
        else:
            raw_content = source.read()
            if isinstance(raw_content, bytes):
                raw_content = raw_content.decode("utf-8")
            filename = getattr(source, "name", "unknown.json")
            file_size = len(raw_content.encode("utf-8"))

        # Parse JSON
        if filename.endswith(".jsonl"):
            # JSON Lines format
            lines = [json.loads(line) for line in raw_content.strip().split("\n") if line.strip()]
            data = lines
        else:
            data = json.loads(raw_content)

        # Extract content
        if self.content_key:
            if isinstance(data, list):
                content = "\n\n".join(
                    str(item.get(self.content_key, item)) for item in data
                )
            else:
                content = str(data.get(self.content_key, data))
        else:
            content = json.dumps(data, indent=2, ensure_ascii=False)

        doc_id = self._generate_id(content)
        meta = metadata or {}

        return Document(
            id=doc_id,
            content=content,
            metadata=DocumentMetadata(
                title=meta.get("title", filename),
                source=str(source) if isinstance(source, (str, Path)) else filename,
                mime_type="application/json",
                file_size=file_size,
                language=meta.get("language"),
                author=meta.get("author"),
                created_at=meta.get("created_at"),
                tags=meta.get("tags", []),
                custom={**(meta.get("custom", {})), "raw_data": data if isinstance(data, dict) else None},
            ),
            status=DocumentStatus.PENDING,
            acl=meta.get("acl", ["*"]),
        )

    def _generate_id(self, content: str) -> str:
        return hashlib.sha256(content.encode()).hexdigest()[:16]


class ExcelLoader(DocumentLoader):
    """Load Excel (.xlsx, .xls) documents."""

    supported_extensions = [".xlsx", ".xls"]
    supported_mimetypes = [
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ]

    def __init__(
        self,
        content_columns: list[str] | None = None,
        sheet_name: str | int | None = 0,
    ):
        """
        Args:
            content_columns: Columns to include in content (None = all)
            sheet_name: Sheet to load (name or index, default first sheet)
        """
        self.content_columns = content_columns
        self.sheet_name = sheet_name

    async def load(
        self,
        source: str | Path | BinaryIO,
        metadata: dict[str, Any] | None = None,
    ) -> Document:
        try:
            import openpyxl
        except ImportError:
            raise ImportError("openpyxl required for Excel loading: pip install openpyxl")

        if isinstance(source, (str, Path)):
            path = Path(source)
            workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
            filename = path.name
            file_size = path.stat().st_size
        else:
            workbook = openpyxl.load_workbook(source, read_only=True, data_only=True)
            filename = getattr(source, "name", "unknown.xlsx")
            source.seek(0)
            file_size = len(source.read())
            source.seek(0)

        # Get sheet
        if isinstance(self.sheet_name, int):
            sheet = workbook.worksheets[self.sheet_name]
        elif self.sheet_name:
            sheet = workbook[self.sheet_name]
        else:
            sheet = workbook.active

        # Read all rows
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            workbook.close()
            return self._create_empty_document(filename, file_size, metadata)

        # First row as headers
        headers = [str(h) if h else f"col_{i}" for i, h in enumerate(rows[0])]
        data_rows = rows[1:]

        # Select columns
        if self.content_columns:
            col_indices = [
                i for i, h in enumerate(headers)
                if h in self.content_columns
            ]
            selected_headers = [headers[i] for i in col_indices]
        else:
            col_indices = list(range(len(headers)))
            selected_headers = headers

        # Build content
        content_parts = []
        for row in data_rows:
            row_values = [str(row[i]) if i < len(row) and row[i] is not None else "" for i in col_indices]
            row_text = " | ".join(f"{h}: {v}" for h, v in zip(selected_headers, row_values))
            content_parts.append(row_text)

        content = "\n".join(content_parts)
        workbook.close()

        doc_id = self._generate_id(content)
        meta = metadata or {}

        return Document(
            id=doc_id,
            content=content,
            metadata=DocumentMetadata(
                title=meta.get("title", filename),
                source=str(source) if isinstance(source, (str, Path)) else filename,
                mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                file_size=file_size,
                language=meta.get("language"),
                author=meta.get("author"),
                created_at=meta.get("created_at"),
                tags=meta.get("tags", []),
                custom={
                    **(meta.get("custom", {})),
                    "columns": headers,
                    "row_count": len(data_rows),
                    "sheet_name": sheet.title,
                },
            ),
            status=DocumentStatus.PENDING,
            acl=meta.get("acl", ["*"]),
        )

    def _create_empty_document(
        self,
        filename: str,
        file_size: int,
        metadata: dict[str, Any] | None,
    ) -> Document:
        """Create document for empty Excel file."""
        meta = metadata or {}
        return Document(
            id=hashlib.sha256(filename.encode()).hexdigest()[:16],
            content="",
            metadata=DocumentMetadata(
                title=meta.get("title", filename),
                source=filename,
                mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                file_size=file_size,
                tags=meta.get("tags", []),
            ),
            status=DocumentStatus.PENDING,
            acl=meta.get("acl", ["*"]),
        )

    def _generate_id(self, content: str) -> str:
        return hashlib.sha256(content.encode()).hexdigest()[:16]


class CSVLoader(DocumentLoader):
    """Load CSV documents."""

    supported_extensions = [".csv", ".tsv"]
    supported_mimetypes = ["text/csv"]

    def __init__(self, content_columns: list[str] | None = None):
        """
        Args:
            content_columns: Columns to include in content (None = all)
        """
        self.content_columns = content_columns

    async def load(
        self,
        source: str | Path | BinaryIO,
        metadata: dict[str, Any] | None = None,
    ) -> Document:
        import csv
        import io

        if isinstance(source, (str, Path)):
            path = Path(source)
            with open(path, "r", encoding="utf-8") as f:
                raw_content = f.read()
            filename = path.name
            file_size = path.stat().st_size
        else:
            raw_content = source.read()
            if isinstance(raw_content, bytes):
                raw_content = raw_content.decode("utf-8")
            filename = getattr(source, "name", "unknown.csv")
            file_size = len(raw_content.encode("utf-8"))

        # Parse CSV
        delimiter = "\t" if filename.endswith(".tsv") else ","
        reader = csv.DictReader(io.StringIO(raw_content), delimiter=delimiter)

        rows = list(reader)
        columns = reader.fieldnames or []

        # Build content
        if self.content_columns:
            selected_cols = [c for c in self.content_columns if c in columns]
        else:
            selected_cols = columns

        content_parts = []
        for row in rows:
            row_text = " | ".join(f"{k}: {row.get(k, '')}" for k in selected_cols)
            content_parts.append(row_text)

        content = "\n".join(content_parts)

        doc_id = self._generate_id(content)
        meta = metadata or {}

        return Document(
            id=doc_id,
            content=content,
            metadata=DocumentMetadata(
                title=meta.get("title", filename),
                source=str(source) if isinstance(source, (str, Path)) else filename,
                mime_type="text/csv",
                file_size=file_size,
                language=meta.get("language"),
                author=meta.get("author"),
                created_at=meta.get("created_at"),
                tags=meta.get("tags", []),
                custom={
                    **(meta.get("custom", {})),
                    "columns": columns,
                    "row_count": len(rows),
                },
            ),
            status=DocumentStatus.PENDING,
            acl=meta.get("acl", ["*"]),
        )

    def _generate_id(self, content: str) -> str:
        return hashlib.sha256(content.encode()).hexdigest()[:16]


# =============================================================================
# Loader Registry
# =============================================================================

LOADERS: dict[str, type[DocumentLoader]] = {
    ".txt": TextLoader,
    ".md": TextLoader,
    ".rst": TextLoader,
    ".log": TextLoader,
    ".pdf": PDFLoader,
    ".html": HTMLLoader,
    ".htm": HTMLLoader,
    ".docx": DocxLoader,
    ".json": JSONLoader,
    ".jsonl": JSONLoader,
    ".csv": CSVLoader,
    ".tsv": CSVLoader,
    ".xlsx": ExcelLoader,
    ".xls": ExcelLoader,
}


async def load_document(
    source: str | Path | BinaryIO,
    metadata: dict[str, Any] | None = None,
    loader_class: type[DocumentLoader] | None = None,
) -> Document:
    """Load a document using appropriate loader.

    Args:
        source: File path, URL, or file-like object
        metadata: Optional metadata to attach
        loader_class: Override automatic loader selection

    Returns:
        Loaded Document object
    """
    if loader_class:
        loader = loader_class()
        return await loader.load(source, metadata)

    # Auto-detect loader based on extension
    if isinstance(source, (str, Path)):
        path = Path(source)
        ext = path.suffix.lower()
    else:
        name = getattr(source, "name", "")
        ext = Path(name).suffix.lower() if name else ""

    if ext not in LOADERS:
        logger.warning("No loader for extension, using TextLoader", ext=ext)
        loader = TextLoader()
    else:
        loader = LOADERS[ext]()

    return await loader.load(source, metadata)
